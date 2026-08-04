"""Deterministically generates all Stage 0 synthetic data for the
resume-job-matcher: 6 synthetic job descriptions and 25+ synthetic
resumes (PDF/DOCX), including the deliberately broken files needed to
exercise the ingestion/validation edge cases in SPECIFICATION.md
Section 8 and Section 19.

Run with:  python sample_data/generate.py

All names, companies, and contact details below are fabricated. No
real personal data is used anywhere in this file or its output.
"""

from __future__ import annotations

import random
import shutil
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

random.seed(42)

BASE_DIR = Path(__file__).resolve().parent
JOBS_DIR = BASE_DIR / "jobs"
RESUMES_DIR = BASE_DIR / "synthetic_resumes"

PAGE_WIDTH, PAGE_HEIGHT = 612, 792  # US letter, points
MARGIN_LEFT, MARGIN_TOP, MARGIN_RIGHT = 50, 50, 50
LINE_HEIGHT = 13
FONT_SIZE = 10
FONT = "helv"


# --------------------------------------------------------------------------
# Low-level file writers
# --------------------------------------------------------------------------

def write_txt(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _wrap_line(line: str, max_chars: int = 100) -> list[str]:
    """Simple word-wrap so long lines don't run off the PDF page."""
    if len(line) <= max_chars:
        return [line]
    words = line.split(" ")
    wrapped, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) > max_chars and current:
            wrapped.append(current)
            current = word
        else:
            current = candidate
    if current:
        wrapped.append(current)
    return wrapped


def write_pdf(path: Path, text: str, encrypt: bool = False, password: str = "") -> None:
    """Renders `text` (newline-separated) into a real, text-layer PDF.

    Lines are drawn explicitly (not via insert_textbox) so nothing is
    ever silently truncated, and pagination is handled manually so
    long resumes span multiple pages deterministically.
    """
    raw_lines: list[str] = []
    for line in text.split("\n"):
        raw_lines.extend(_wrap_line(line) if line else [""])

    doc = fitz.open()
    usable_height = PAGE_HEIGHT - MARGIN_TOP - MARGIN_LEFT
    lines_per_page = max(1, usable_height // LINE_HEIGHT)

    for start in range(0, len(raw_lines), lines_per_page):
        page = doc.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
        y = MARGIN_TOP
        for line in raw_lines[start:start + lines_per_page]:
            if line:
                page.insert_text((MARGIN_LEFT, y), line, fontsize=FONT_SIZE, fontname=FONT)
            y += LINE_HEIGHT

    if encrypt:
        doc.save(
            str(path),
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw=password or "owner-secret",
            user_pw=password or "user-secret",
        )
    else:
        doc.save(str(path))
    doc.close()


HEADING_LINES = {
    "SUMMARY", "PROFESSIONAL SUMMARY",
    "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY",
    "TECHNICAL SKILLS", "SKILLS", "CORE COMPETENCIES",
    "EDUCATION",
    "CERTIFICATIONS", "LICENSES & CERTIFICATIONS",
    "PROJECTS",
}


def write_docx(path: Path, text: str, table_section: str | None = None) -> None:
    """Renders `text` into a real .docx file.

    If `table_section` is set (e.g. "TECHNICAL SKILLS"), that section's
    content is emitted as a Word table instead of paragraphs, to
    exercise the table-cell extraction path in Section 8.2.
    """
    document = Document()
    lines = text.split("\n")
    in_table_section = False
    table_rows: list[str] = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            table = document.add_table(rows=0, cols=1)
            for row_text in table_rows:
                cells = table.add_row().cells
                cells[0].text = row_text
            table_rows = []

    for line in lines:
        stripped = line.strip()
        is_heading = stripped.upper() in HEADING_LINES and stripped != ""

        if is_heading:
            flush_table()
            in_table_section = table_section is not None and stripped.upper() == table_section.upper()
            document.add_heading(stripped, level=2)
            continue

        if in_table_section and stripped:
            table_rows.append(stripped)
            continue

        if not stripped:
            flush_table()
            document.add_paragraph("")
            continue

        document.add_paragraph(stripped)

    flush_table()
    document.save(str(path))


# --------------------------------------------------------------------------
# Job descriptions (6): varied heading styles per SPECIFICATION.md 10.2
# --------------------------------------------------------------------------

JOBS = [
    {
        "filename": "job_01_data_analyst_standard.txt",
        "title": "Data Analyst",
        "text": """Data Analyst

About Us
Northfield Retail Analytics helps mid-size retailers understand their customers through data.

Responsibilities
- Build and maintain recurring dashboards in Power BI for merchandising and marketing stakeholders.
- Write SQL queries against the retail data warehouse to answer ad hoc business questions.
- Clean and validate incoming sales and inventory datasets before they reach reporting layers.
- Present findings and recommendations to non-technical stakeholders in weekly review meetings.
- Partner with the marketing and merchandising teams to define new reporting requirements.

Requirements
- Must have strong SQL skills, including joins, aggregations, and window functions.
- Must have advanced Microsoft Excel skills, including pivot tables and XLOOKUP.
- A bachelor's degree in Data Analytics, Statistics, Economics, or a related field is required.
- Candidates must have at least 2 years of experience in a data analysis or reporting role.

Preferred Qualifications
- Experience with Power BI is preferred.
- Familiarity with Python for data cleaning is a plus.
- Exposure to Tableau is a bonus.

Benefits
Health insurance, 401k matching, and unlimited PTO.
""",
    },
    {
        "filename": "job_02_data_analyst_altheadings.txt",
        "title": "Data Analyst II",
        "text": """Data Analyst II

Who We Are
Bridgeline Health Systems operates community clinics across the region.

What You Will Do
- Maintain weekly and monthly operational reports for clinic administrators.
- Query the electronic health records data mart using SQL to extract utilization metrics.
- Build Excel-based summary workbooks for the finance team.
- Investigate data quality issues in incoming reports and document root causes.

What You Need
- SQL is required for this role; you will write queries daily.
- Advanced Excel skills are mandatory, including pivot tables and lookup formulas.
- A bachelor's degree in Statistics, Public Health, or a quantitative field is required.
- This role requires a minimum of 3 years of experience in healthcare or operations reporting.

Compensation
Salary range disclosed upon interview.
""",
    },
    {
        "filename": "job_03_data_engineer.txt",
        "title": "Data Engineer",
        "text": """Data Engineer

About the Team
We are a five-person platform engineering team supporting the whole analytics org.

Minimum Qualifications
- Strong SQL skills are required, including query optimization.
- Python is required for building and maintaining ETL pipelines.
- Experience designing and building data pipelines is essential.
- A bachelor's degree in Computer Science or a related field is required.
- Candidates must have at least 3 years of experience in a data engineering role.

Key Activities
- Design, build, and maintain ETL pipelines that move data from source systems into the warehouse.
- Monitor pipeline health and resolve data pipeline failures.
- Optimize SQL queries and warehouse tables for performance.
- Collaborate with analysts and data scientists on data model design.
- Document data lineage for critical tables.

Nice to Have
- Experience with Apache Airflow is a plus.
- Familiarity with AWS, including S3 and Redshift, is preferred.
- Exposure to Docker is a bonus.
""",
    },
    {
        "filename": "job_04_bi_analyst.txt",
        "title": "BI Analyst",
        "text": """BI Analyst

The Role
- Design and maintain BI dashboards used across the leadership team.
- Define and track key performance indicators for the business.
- Write SQL against the analytics warehouse to build new reports.
- Maintain Excel-based ad hoc reporting for the finance team.

Must Have
- Power BI experience is required; it is the team's primary reporting tool.
- SQL is required for building and troubleshooting reports.
- Advanced Excel skills, including pivot tables, are required.
- A bachelor's degree is required.

Bonus Points
- Experience with Tableau is a plus.
- Familiarity with Python is preferred.

Perks
Free lunch on Fridays and a hybrid schedule.
""",
    },
    {
        "filename": "job_05_data_scientist.txt",
        "title": "Data Scientist",
        "text": """Data Scientist

Requirements
- Python is required for model development.
- SQL is required to query training data from the warehouse.
- Machine learning modeling experience, including classification and regression, is required.
- A master's degree in a quantitative field such as Statistics, Computer Science, or Data Science,
  or equivalent experience of 4 years, is required.
- Candidates must have at least 2 years of relevant experience.

Duties
- Build and validate predictive models for customer churn and demand forecasting.
- Query and prepare training datasets using SQL.
- Communicate model performance and limitations to business stakeholders.
- Monitor deployed models for drift and retrain as needed.

Desired
- Experience with deep learning frameworks is a plus.
- Familiarity with Power BI or Tableau for communicating results is preferred.
- Exposure to AWS SageMaker is a bonus.
""",
    },
    {
        "filename": "job_06_software_engineer.txt",
        "title": "Software Engineer",
        "text": """Software Engineer

What You Need
- Python is required for backend service development.
- SQL is required for working with the application database.
- A bachelor's degree in Computer Science or a related field is required.
- Candidates must have at least 2 years of professional software engineering experience.

Responsibilities
- Build and maintain backend services in Python.
- Write and optimize SQL queries against the application database.
- Participate in code review and on-call rotations.
- Collaborate with product managers to scope new features.

A Plus
- PMP certification is a plus for engineers interested in technical leadership.
- Experience with AWS is preferred.
- Familiarity with Docker is a bonus.
""",
    },
]


# --------------------------------------------------------------------------
# Resume text builder
# --------------------------------------------------------------------------

def build_resume_text(cand: dict) -> str:
    lines: list[str] = [cand["name"], f"{cand['email']} | {cand['phone']} | {cand['location']}", ""]

    if cand.get("flat_text"):
        # NO_HEADINGS case: whole body is one unsectioned block.
        lines.append(cand["flat_text"])
        return "\n".join(lines)

    if cand.get("summary"):
        lines += ["SUMMARY", cand["summary"], ""]

    lines.append("PROFESSIONAL EXPERIENCE")
    for job in cand.get("experience", []):
        header = f"{job['title']} - {job['company']}"
        if job.get("dates"):
            header += f" ({job['dates']})"
        lines.append(header)
        for bullet in job.get("bullets", []):
            lines.append(f"- {bullet}")
        lines.append("")

    lines.append("TECHNICAL SKILLS")
    lines.append(", ".join(cand.get("skills", [])))
    lines.append("")

    lines.append("EDUCATION")
    lines.append(cand.get("education", "Coursework in progress; no degree conferred."))

    if cand.get("certifications"):
        lines += ["", "CERTIFICATIONS"]
        for cert in cand["certifications"]:
            lines.append(f"- {cert}")

    if cand.get("interests"):
        lines += ["", "INTERESTS", cand["interests"]]

    return "\n".join(lines)


# --------------------------------------------------------------------------
# Candidates (20 "normal" resumes, targeted at the 6 jobs above)
# --------------------------------------------------------------------------

CANDIDATES = [
    # ---- Job 1: Data Analyst (standard) ----------------------------------
    {
        "filename": "job1_strong_match_analyst",
        "format": "pdf",
        "case": "strong_match",
        "target_job": "job_01_data_analyst_standard",
        "name": "Jordan Ellis",
        "email": "jordan.ellis.demo@example.com",
        "phone": "(555) 010-1001",
        "location": "Columbus, OH",
        "summary": "Data analyst with 3 years of experience building dashboards and reporting pipelines "
                    "for retail merchandising teams using SQL, Power BI, and Python.",
        "experience": [
            {
                "title": "Data Analyst",
                "company": "Harborview Retail Group",
                "dates": "Jun 2022 - Present",
                "bullets": [
                    "Built and maintained recurring Power BI dashboards for merchandising and marketing stakeholders.",
                    "Wrote SQL queries against the retail data warehouse, including joins and window functions, to answer ad hoc business questions.",
                    "Cleaned and validated incoming sales and inventory datasets before they reached reporting layers.",
                    "Presented findings and recommendations to non-technical stakeholders in weekly review meetings.",
                    "Used Python scripts to automate recurring data cleaning steps.",
                ],
            },
            {
                "title": "Reporting Analyst",
                "company": "Cedar Grove Supply Co.",
                "dates": "Jul 2021 - May 2022",
                "bullets": [
                    "Built Excel pivot table reports for regional sales managers.",
                    "Queried operational data with SQL to track inventory turnover.",
                ],
            },
        ],
        "skills": ["SQL", "Excel", "Power BI", "Python", "Data Cleaning"],
        "education": "Bachelor of Science in Data Analytics, State University, completed.",
    },
    {
        "filename": "job1_good_gaps_analyst",
        "format": "docx",
        "case": "good_but_gaps",
        "target_job": "job_01_data_analyst_standard",
        "name": "Morgan Patel",
        "email": "morgan.patel.demo@example.com",
        "phone": "(555) 010-1002",
        "location": "Austin, TX",
        "summary": "Analyst supporting merchandising operations with SQL reporting and Excel modeling.",
        "experience": [
            {
                "title": "Business Analyst",
                "company": "Prairie Home Goods",
                "dates": "Mar 2023 - Present",
                "bullets": [
                    "Wrote SQL queries to pull weekly sales performance data for category managers.",
                    "Built Excel workbooks with pivot tables and XLOOKUP formulas for quarterly business reviews.",
                    "Used Power BI to publish a self-service dashboard for the merchandising team.",
                ],
            },
            {
                "title": "Junior Analyst",
                "company": "Prairie Home Goods",
                "dates": "Jan 2022 - Feb 2023",
                "bullets": [
                    "Supported senior analysts with data pulls and report formatting.",
                ],
            },
        ],
        "skills": ["SQL", "Excel", "Power BI"],
        "education": "Bachelor of Business Administration, Concentration in Business Analytics, Lakeshore College, completed.",
    },
    {
        "filename": "job1_keyword_stuffer_analyst",
        "format": "pdf",
        "case": "keyword_stuffer",
        "target_job": "job_01_data_analyst_standard",
        "name": "Casey Nguyen",
        "email": "casey.nguyen.demo@example.com",
        "phone": "(555) 010-1003",
        "location": "Denver, CO",
        "summary": "Results-driven professional seeking a data analyst opportunity.",
        "experience": [
            {
                "title": "Administrative Coordinator",
                "company": "Summit Office Solutions",
                "dates": "May 2021 - Present",
                "bullets": [
                    "Responsible for daily administrative tasks and supporting the front office team.",
                    "Assisted with scheduling and general correspondence as needed.",
                    "Maintained office supply inventory and vendor relationships.",
                ],
            },
        ],
        "skills": ["SQL", "Excel", "Power BI", "Python", "Tableau", "Data Analysis", "Data Visualization"],
        "education": "Bachelor of Arts in Communications, Ridgeline University, completed.",
    },
    {
        "filename": "job1_career_changer_analyst",
        "format": "docx",
        "case": "career_changer",
        "target_job": "job_01_data_analyst_standard",
        "name": "Riley Thompson",
        "email": "riley.thompson.demo@example.com",
        "phone": "(555) 010-1004",
        "location": "Madison, WI",
        "summary": "Former educator transitioning into data analysis, self-taught in SQL and Excel.",
        "experience": [
            {
                "title": "High School Mathematics Teacher",
                "company": "Fairview High School",
                "dates": "Aug 2018 - Jun 2024",
                "bullets": [
                    "Planned and delivered daily mathematics lessons for grades 9 through 12.",
                    "Tracked student performance and communicated progress to parents.",
                    "Organized department budgeting spreadsheets each semester.",
                ],
            },
        ],
        "skills": ["SQL (self-taught)", "Excel"],
        "education": "Bachelor of Arts in Education, Fairview State College, completed.",
    },
    {
        "filename": "job1_missing_preferred_skill_analyst",
        "format": "pdf",
        "case": "missing_one_preferred",
        "target_job": "job_01_data_analyst_standard",
        "name": "Dakota Reyes",
        "email": "dakota.reyes.demo@example.com",
        "phone": "(555) 010-1005",
        "location": "Phoenix, AZ",
        "summary": "Data analyst with 4 years of experience in retail reporting and dashboarding, "
                    "strong in SQL, Excel, and Power BI.",
        "experience": [
            {
                "title": "Senior Data Analyst",
                "company": "Sagebrush Retail Analytics",
                "dates": "Feb 2021 - Present",
                "bullets": [
                    "Built and maintained recurring Power BI dashboards for merchandising and marketing stakeholders.",
                    "Wrote complex SQL queries with joins and window functions against the retail warehouse.",
                    "Cleaned and validated incoming sales and inventory datasets before reporting.",
                    "Automated recurring reporting tasks using Python scripts.",
                ],
            },
        ],
        "skills": ["SQL", "Excel", "Power BI", "Python"],
        "education": "Bachelor of Science in Statistics, Meridian University, completed.",
    },

    # ---- Job 2: Data Analyst II (alt headings, no preferred) -------------
    {
        "filename": "job2_strong_match_analyst2",
        "format": "pdf",
        "case": "strong_match",
        "target_job": "job_02_data_analyst_altheadings",
        "name": "Avery Collins",
        "email": "avery.collins.demo@example.com",
        "phone": "(555) 010-1006",
        "location": "Cleveland, OH",
        "summary": "Healthcare operations analyst with 4 years of SQL and Excel reporting experience.",
        "experience": [
            {
                "title": "Operations Analyst",
                "company": "Bridgeline Health Systems",
                "dates": "Sep 2021 - Present",
                "bullets": [
                    "Maintained weekly and monthly operational reports for clinic administrators.",
                    "Queried the electronic health records data mart using SQL to extract utilization metrics.",
                    "Built Excel-based summary workbooks for the finance team.",
                    "Investigated data quality issues in incoming reports and documented root causes.",
                ],
            },
        ],
        "skills": ["SQL", "Excel", "Healthcare Reporting"],
        "education": "Bachelor of Science in Public Health, Coastal University, completed.",
    },
    {
        "filename": "job2_good_gaps_analyst2",
        "format": "docx",
        "case": "at_minimum",
        "target_job": "job_02_data_analyst_altheadings",
        "name": "Sam Whitfield",
        "email": "sam.whitfield.demo@example.com",
        "phone": "(555) 010-1007",
        "location": "Pittsburgh, PA",
        "summary": "Operations reporting analyst with SQL experience in clinic settings.",
        "experience": [
            {
                "title": "Reporting Coordinator",
                "company": "Elmwood Community Clinics",
                "dates": "Jun 2021 - Present",
                "bullets": [
                    "Queried the clinic data mart using SQL to support monthly utilization reports.",
                    "Investigated recurring data quality issues in patient scheduling exports.",
                ],
            },
        ],
        "skills": ["SQL", "Excel"],
        "education": "Bachelor of Science in Statistics, Ironwood University, completed.",
    },

    # ---- Job 3: Data Engineer ---------------------------------------------
    {
        "filename": "job3_strong_match_engineer",
        "format": "pdf",
        "case": "strong_match",
        "target_job": "job_03_data_engineer",
        "name": "Taylor Brooks",
        "email": "taylor.brooks.demo@example.com",
        "phone": "(555) 010-1008",
        "location": "Seattle, WA",
        "summary": "Data engineer with 4 years of experience building ETL pipelines on AWS.",
        "experience": [
            {
                "title": "Data Engineer",
                "company": "Cascade Freight Analytics",
                "dates": "Jan 2021 - Present",
                "bullets": [
                    "Designed, built, and maintained ETL pipelines that move data from source systems into the warehouse.",
                    "Monitored pipeline health and resolved data pipeline failures using Airflow.",
                    "Optimized SQL queries and warehouse tables for performance.",
                    "Built pipelines on AWS using S3 and Redshift, containerized with Docker.",
                    "Documented data lineage for critical tables.",
                ],
            },
        ],
        "skills": ["SQL", "Python", "Airflow", "AWS", "Docker", "ETL"],
        "education": "Bachelor of Science in Computer Science, Rainier Institute of Technology, completed.",
    },
    {
        "filename": "job3_good_gaps_engineer_table",
        "format": "docx",
        "case": "good_but_gaps_table_based",
        "target_job": "job_03_data_engineer",
        "name": "Jamie Ortiz",
        "email": "jamie.ortiz.demo@example.com",
        "phone": "(555) 010-1009",
        "location": "Portland, OR",
        "summary": "Engineer focused on building and maintaining data pipelines for analytics teams.",
        "experience": [
            {
                "title": "Data Pipeline Engineer",
                "company": "Willamette Analytics",
                "dates": "Apr 2022 - Present",
                "bullets": [
                    "Built and maintained data pipelines moving data from operational systems into the warehouse.",
                    "Wrote SQL to optimize slow-running warehouse queries.",
                    "Used Python to write ingestion scripts for new source systems.",
                ],
            },
        ],
        "skills": ["SQL", "Python", "Data Pipelines"],
        "education": "Bachelor of Science in Information Systems, Willamette State University, completed.",
        "table_section": "TECHNICAL SKILLS",
    },
    {
        "filename": "job3_no_headings_engineer",
        "format": "pdf",
        "case": "no_headings",
        "target_job": "job_03_data_engineer",
        "name": "Charlie Fenwick",
        "email": "charlie.fenwick.demo@example.com",
        "phone": "(555) 010-1010",
        "location": "San Jose, CA",
        "flat_text": (
            "Data engineer with three years of experience. Designed, built, and maintained ETL pipelines "
            "that move data from source systems into the warehouse at Bayline Data Co. from March 2022 to "
            "present. Wrote and optimized SQL queries against the warehouse daily. Used Python and Airflow "
            "to schedule and monitor pipeline runs. Worked with AWS S3 and Redshift for storage and "
            "compute. Bachelor of Science in Computer Science, Bayline Technical University, completed."
        ),
    },
    {
        "filename": "job3_missing_dates_engineer",
        "format": "docx",
        "case": "missing_dates",
        "target_job": "job_03_data_engineer",
        "name": "Skyler Vance",
        "email": "skyler.vance.demo@example.com",
        "phone": "(555) 010-1011",
        "location": "Chicago, IL",
        "summary": "Data engineer experienced with SQL and Python pipeline development.",
        "experience": [
            {
                "title": "Data Engineer",
                "company": "Lakeside Systems",
                "dates": "",
                "bullets": [
                    "Designed and maintained ETL pipelines moving data into the analytics warehouse.",
                    "Optimized SQL queries against the warehouse for performance.",
                    "Built ingestion jobs in Python for third-party data sources.",
                ],
            },
        ],
        "skills": ["SQL", "Python", "ETL"],
        "education": "Bachelor of Science in Computer Science, Lakeside University, completed.",
    },
    {
        "filename": "job3_year_only_dates_engineer",
        "format": "pdf",
        "case": "year_only_dates",
        "target_job": "job_03_data_engineer",
        "name": "Peyton Marsh",
        "email": "peyton.marsh.demo@example.com",
        "phone": "(555) 010-1012",
        "location": "Boston, MA",
        "summary": "Data engineer with pipeline development experience in SQL and Python.",
        "experience": [
            {
                "title": "Data Engineer",
                "company": "Beacon Hill Data Co.",
                "dates": "2019 - 2022",
                "bullets": [
                    "Designed and built ETL pipelines to move data into the central warehouse.",
                    "Wrote SQL to optimize reporting tables for downstream analysts.",
                    "Used Python for scripting recurring ingestion jobs.",
                ],
            },
        ],
        "skills": ["SQL", "Python", "ETL"],
        "education": "Bachelor of Science in Computer Science, Beacon Hill College, completed.",
    },

    # ---- Job 4: BI Analyst (no experience minimum) ------------------------
    {
        "filename": "job4_strong_match_bi",
        "format": "docx",
        "case": "strong_match",
        "target_job": "job_04_bi_analyst",
        "name": "Rowan Ahmed",
        "email": "rowan.ahmed.demo@example.com",
        "phone": "(555) 010-1013",
        "location": "Minneapolis, MN",
        "summary": "BI analyst building executive dashboards and KPI tracking in Power BI and Tableau.",
        "experience": [
            {
                "title": "BI Analyst",
                "company": "Northstar Manufacturing",
                "dates": "May 2021 - Present",
                "bullets": [
                    "Designed and maintained BI dashboards used across the leadership team in Power BI.",
                    "Defined and tracked key performance indicators for the business.",
                    "Wrote SQL against the analytics warehouse to build new reports.",
                    "Maintained Excel-based ad hoc reporting for the finance team.",
                    "Built a supplemental reporting layer in Tableau for the sales team, with light Python scripting.",
                ],
            },
        ],
        "skills": ["Power BI", "SQL", "Excel", "Tableau", "Python"],
        "education": "Bachelor of Science in Business Analytics, Northstar University, completed.",
    },
    {
        "filename": "job4_good_gaps_bi",
        "format": "pdf",
        "case": "good_but_gaps",
        "target_job": "job_04_bi_analyst",
        "name": "Emerson Blake",
        "email": "emerson.blake.demo@example.com",
        "phone": "(555) 010-1014",
        "location": "Kansas City, MO",
        "summary": "Reporting analyst focused on Power BI dashboards and SQL reporting.",
        "experience": [
            {
                "title": "Reporting Analyst",
                "company": "Heartland Manufacturing",
                "dates": "Feb 2022 - Present",
                "bullets": [
                    "Built and maintained Power BI dashboards for the operations leadership team.",
                    "Wrote SQL queries against the analytics warehouse for new reporting requests.",
                ],
            },
        ],
        "skills": ["Power BI", "SQL", "Excel"],
        "education": "Bachelor of Science in Finance, Heartland College, completed.",
    },

    # ---- Job 5: Data Scientist --------------------------------------------
    {
        "filename": "job5_strong_match_scientist",
        "format": "docx",
        "case": "strong_match",
        "target_job": "job_05_data_scientist",
        "name": "Quinn Delgado",
        "email": "quinn.delgado.demo@example.com",
        "phone": "(555) 010-1015",
        "location": "San Francisco, CA",
        "summary": "Data scientist with 3 years of experience building machine learning models for churn and demand forecasting.",
        "experience": [
            {
                "title": "Data Scientist",
                "company": "Bayfront Analytics",
                "dates": "Jul 2021 - Present",
                "bullets": [
                    "Built and validated machine learning models for customer churn using classification techniques.",
                    "Built regression models for demand forecasting.",
                    "Queried and prepared training datasets using SQL.",
                    "Communicated model performance and limitations to business stakeholders using Tableau and Power BI.",
                    "Wrote model training and evaluation code in Python.",
                ],
            },
        ],
        "skills": ["Python", "SQL", "Machine Learning", "Tableau", "Power BI"],
        "education": "Master of Science in Data Science, Bayfront University, completed.",
    },
    {
        "filename": "job5_degree_or_equivalent_scientist",
        "format": "pdf",
        "case": "degree_or_equivalent",
        "target_job": "job_05_data_scientist",
        "name": "Harper Nakamura",
        "email": "harper.nakamura.demo@example.com",
        "phone": "(555) 010-1016",
        "location": "Chicago, IL",
        "summary": "Data scientist, no formal graduate degree, with 5 years of hands-on machine learning experience.",
        "experience": [
            {
                "title": "Data Scientist",
                "company": "Millbrook Analytics",
                "dates": "Jun 2019 - Present",
                "bullets": [
                    "Built and validated classification and regression models for pricing optimization.",
                    "Queried and prepared training datasets using SQL against the analytics warehouse.",
                    "Wrote Python code for model training, evaluation, and monitoring.",
                    "Communicated model performance to business stakeholders.",
                ],
            },
        ],
        "skills": ["Python", "SQL", "Machine Learning"],
        "education": "Some college coursework in mathematics; no degree conferred.",
    },
    {
        "filename": "job5_good_gaps_scientist",
        "format": "pdf",
        "case": "good_but_gaps",
        "target_job": "job_05_data_scientist",
        "name": "Micah Solano",
        "email": "micah.solano.demo@example.com",
        "phone": "(555) 010-1017",
        "location": "Raleigh, NC",
        "summary": "Data scientist with SQL and Python experience in a statistics graduate program.",
        "experience": [
            {
                "title": "Data Scientist",
                "company": "Triangle Insights",
                "dates": "Sep 2022 - Present",
                "bullets": [
                    "Queried training data from the warehouse using SQL for internal analytics projects.",
                    "Wrote Python scripts for data cleaning and feature engineering.",
                ],
            },
        ],
        "skills": ["Python", "SQL", "Machine Learning"],
        "education": "Master of Science in Statistics, Triangle State University, completed.",
    },
    {
        "filename": "job5_weak_scientist",
        "format": "docx",
        "case": "weak_match",
        "target_job": "job_05_data_scientist",
        "name": "Elliot Marsh",
        "email": "elliot.marsh.demo@example.com",
        "phone": "(555) 010-1018",
        "location": "Tampa, FL",
        "summary": "Analyst with SQL reporting experience looking to move into data science.",
        "experience": [
            {
                "title": "Reporting Analyst",
                "company": "Palmview Retail",
                "dates": "Jan 2023 - Present",
                "bullets": [
                    "Queried sales data using SQL for weekly performance reports.",
                    "Built Excel summaries for the merchandising team.",
                ],
            },
        ],
        "skills": ["SQL", "Excel"],
        "education": "Bachelor of Arts in Economics, Palmview College, completed.",
    },

    # ---- Job 6: Software Engineer (PMP preferred) --------------------------
    {
        "filename": "job6_strong_match_swe",
        "format": "docx",
        "case": "strong_match",
        "target_job": "job_06_software_engineer",
        "name": "Reese Chandler",
        "email": "reese.chandler.demo@example.com",
        "phone": "(555) 010-1019",
        "location": "Atlanta, GA",
        "summary": "Backend software engineer with 3 years of experience building Python services.",
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Peachtree Software",
                "dates": "Aug 2021 - Present",
                "bullets": [
                    "Built and maintained backend services in Python.",
                    "Wrote and optimized SQL queries against the application database.",
                    "Participated in code review and on-call rotations.",
                    "Collaborated with product managers to scope new features.",
                    "Deployed services to AWS using Docker containers.",
                ],
            },
        ],
        "skills": ["Python", "SQL", "AWS", "Docker"],
        "education": "Bachelor of Science in Computer Science, Peachtree Institute, completed.",
        "certifications": ["Project Management Professional (PMP), certified 2023"],
    },
    {
        "filename": "job6_pmp_candidate_engineer",
        "format": "pdf",
        "case": "pmp_candidate",
        "target_job": "job_06_software_engineer",
        "name": "Finley Osei",
        "email": "finley.osei.demo@example.com",
        "phone": "(555) 010-1020",
        "location": "Charlotte, NC",
        "summary": "Software engineer with backend Python experience, pursuing PMP certification.",
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Uptown Software Labs",
                "dates": "Mar 2022 - Present",
                "bullets": [
                    "Built and maintained backend services in Python.",
                    "Wrote SQL queries against the application database.",
                    "Participated in on-call rotations and code review.",
                ],
            },
        ],
        "skills": ["Python", "SQL"],
        "education": "Bachelor of Science in Computer Science, Uptown University, completed.",
        "certifications": ["PMP candidate, exam scheduled for late 2026"],
    },
]


# --------------------------------------------------------------------------
# Generation
# --------------------------------------------------------------------------

def generate_jobs() -> None:
    JOBS_DIR.mkdir(parents=True, exist_ok=True)
    for job in JOBS:
        write_txt(JOBS_DIR / job["filename"], job["text"])
        print(f"  wrote {job['filename']}")


def generate_candidate_resumes() -> None:
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    for cand in CANDIDATES:
        text = build_resume_text(cand)
        stem = cand["filename"]
        if cand["format"] == "pdf":
            write_pdf(RESUMES_DIR / f"{stem}.pdf", text)
        else:
            write_docx(RESUMES_DIR / f"{stem}.docx", text, table_section=cand.get("table_section"))
        print(f"  wrote {stem}.{cand['format']}  [{cand['case']}]")


def generate_edge_cases() -> None:
    """Deliberately broken / boundary files exercising Section 8 validation."""

    # 1. Exact duplicate of an already-accepted PDF (byte-identical, new name)
    #    -> must be detected by SHA-256, not filename.
    source = RESUMES_DIR / "job1_strong_match_analyst.pdf"
    duplicate = RESUMES_DIR / "job1_strong_match_analyst_duplicate.pdf"
    shutil.copyfile(source, duplicate)
    print(f"  wrote {duplicate.name}  [exact_duplicate]")

    # 2. Corrupt file: correct extension, but not a valid PDF structure at all.
    corrupt = RESUMES_DIR / "edge_corrupt_file.pdf"
    corrupt.write_bytes(b"%PDF-1.4\nThis is not a real PDF body, just garbage bytes.\n%%EOF-TRUNCATED")
    print(f"  wrote {corrupt.name}  [corrupt]")

    # 3. Password-protected PDF: valid content, but encrypted -> CorruptFileError path.
    write_pdf(
        RESUMES_DIR / "edge_password_protected.pdf",
        "Avery Stone\nStrong candidate resume text locked behind a password.\n\n"
        "PROFESSIONAL EXPERIENCE\nData Analyst - Locked Corp (2020 - Present)\n"
        "- This content should never be extracted without the password.\n",
        encrypt=True,
        password="synthetic-test-password",
    )
    print("  wrote edge_password_protected.pdf  [password_protected]")

    # 4. Near-empty "probable scan" PDF: valid PDF, extractable text far below
    #    the 200-character min_extracted_chars threshold.
    write_pdf(RESUMES_DIR / "edge_probable_scan.pdf", "Resume")
    print("  wrote edge_probable_scan.pdf  [probable_scan]")

    # 5. Plain text saved with a .pdf extension: extension lies about content,
    #    magic-bytes signature check must catch this (starts with neither
    #    %PDF nor PK\x03\x04).
    renamed = RESUMES_DIR / "edge_renamed_txt_as_pdf.pdf"
    renamed.write_text(
        "Jordan Casey\nThis file is actually plain text saved with a .pdf extension.\n",
        encoding="utf-8",
    )
    print(f"  wrote {renamed.name}  [signature_mismatch]")

    # 6. Unsupported extension entirely (legacy .doc, not accepted by the MVP).
    unsupported = RESUMES_DIR / "edge_unsupported_filetype.doc"
    unsupported.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1FAKE LEGACY DOC BYTES")
    print(f"  wrote {unsupported.name}  [unsupported_extension]")


def main() -> None:
    print("Generating job descriptions...")
    generate_jobs()
    print("Generating candidate resumes...")
    generate_candidate_resumes()
    print("Generating ingestion edge-case files...")
    generate_edge_cases()

    n_jobs = len(list(JOBS_DIR.glob("*.txt")))
    n_resumes = len([p for p in RESUMES_DIR.iterdir() if p.is_file()])
    print(f"\nDone. {n_jobs} job descriptions, {n_resumes} resume files.")


if __name__ == "__main__":
    main()
