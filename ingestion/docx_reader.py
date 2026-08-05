"""DOCX text extraction (SPECIFICATION.md Section 8.2). Combines
paragraph text and table-cell text into one stream, since a resume's
skills or experience can legitimately be laid out as a table (see
sample_data's table-based synthetic resume).
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from domain.exceptions import CorruptFileError


def extract_docx_text(file_bytes: bytes) -> str:
    """Exact contract from Section 8.2: raises CorruptFileError for any
    failure to read the file (not a real DOCX / ZIP, missing the Word
    document part, python-docx internal error, etc.)."""
    try:
        document = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        table_text = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_text)
    except CorruptFileError:
        raise
    except Exception as exc:
        raise CorruptFileError(f"Unable to extract DOCX text: {exc}") from exc
